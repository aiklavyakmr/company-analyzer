import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import openai
import tempfile
import os

# Set your API key from environment variable
openai.api_key = os.getenv("OPENAI_API_KEY")

STANDARD_PROMPT = "Analyze the following document from the target company and provide a 1–2 page summary focusing on key strategic, financial, and operational insights."

def extract_text(file):
    if file.type == "application/pdf":
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return None

def get_summary(text):
    messages = [
        {"role": "system", "content": "You are an expert in analyzing company documents."},
        {"role": "user", "content": f"{STANDARD_PROMPT}\n\n{text}"}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=messages,
        temperature=0.5,
        max_tokens=1500
    )
    return response['choices'][0]['message']['content']

def save_to_word(summary):
    doc = Document()
    doc.add_heading("Company Summary", level=1)
    for para in summary.split("\n"):
        doc.add_paragraph(para)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

st.title("📊 Company Document Analyzer")

uploaded_file = st.file_uploader("Upload a company document (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("Processing document..."):
        text = extract_text(uploaded_file)
        if text:
            summary = get_summary(text)
            st.subheader("📝 Summary")
            st.text_area("Generated Summary", value=summary, height=300)

            docx_path = save_to_word(summary)
            with open(docx_path, "rb") as file:
                st.download_button(
                    label="📥 Download Summary as Word Document",
                    data=file,
                    file_name="company_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Unsupported file type.")
